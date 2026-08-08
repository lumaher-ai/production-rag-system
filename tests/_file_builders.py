"""In-memory sample-file builders shared across ingestion tests."""

from io import BytesIO

from docx import Document as DocxDocument

_FILLER = (
    "The quarterly figures below summarize regional performance and the "
    "adjustments applied to the prior period. Values are reported in local "
    "currency and rounded to the nearest unit. "
)


def _pad(text: str, target: int) -> str:
    """Bring a page up to the character density of real prose."""
    if not text or target <= len(text):
        return text
    filler = _FILLER * (1 + (target - len(text)) // len(_FILLER))
    return f"{text} {filler}"[:target]


def make_pdf(page_texts: list[str], *, chars_per_page: int = 400) -> bytes:
    """Build a minimal, valid multi-page PDF with literal text per page.

    Hand-rolled (rather than via a rendering lib) so pypdf can extract the exact
    text back out and so page numbering is deterministic for provenance tests.

    Non-empty pages are padded with filler to ``chars_per_page``. A real page of
    prose runs 1,500-3,000 characters, and the ingestion quality gate rejects
    PDFs below 50 — so an unpadded fixture carrying "Page one text." is, by that
    measure, indistinguishable from a scan. Padding keeps fixtures standing in
    for *normal* documents; pass ``chars_per_page=0`` when sparseness is the
    point, or use ``make_scanned_pdf``.

    Empty strings stay empty, which is how a page that yielded no text is built.
    """
    page_texts = [_pad(text, chars_per_page) for text in page_texts]
    n_pages = len(page_texts)
    page_obj_start = 3
    content_obj_start = page_obj_start + n_pages
    font_obj = content_obj_start + n_pages

    objects: list[tuple[int, str]] = []
    kids = " ".join(f"{page_obj_start + i} 0 R" for i in range(n_pages))
    objects.append((1, "<< /Type /Catalog /Pages 2 0 R >>"))
    objects.append((2, f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>"))
    for i in range(n_pages):
        objects.append(
            (
                page_obj_start + i,
                "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Contents {content_obj_start + i} 0 R "
                f"/Resources << /Font << /F1 {font_obj} 0 R >> >> >>",
            )
        )
    for i, text in enumerate(page_texts):
        safe = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        stream = f"BT /F1 24 Tf 72 700 Td ({safe}) Tj ET"
        objects.append(
            (content_obj_start + i, f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream")
        )
    objects.append((font_obj, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))

    body = "%PDF-1.4\n"
    offsets: dict[int, int] = {}
    for num, content in objects:
        offsets[num] = len(body.encode("latin-1"))
        body += f"{num} 0 obj\n{content}\nendobj\n"

    xref_offset = len(body.encode("latin-1"))
    total = len(objects) + 1  # include free object 0
    body += f"xref\n0 {total}\n0000000000 65535 f \n"
    for num in range(1, total):
        body += f"{offsets[num]:010d} 00000 n \n"
    body += f"trailer\n<< /Size {total} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF"
    return body.encode("latin-1")


def make_scanned_pdf(page_count: int, *, text_pages: int = 0) -> bytes:
    """A PDF that extracts to almost nothing — what a scan looks like to pypdf.

    ``text_pages`` gives the first N pages real text and leaves the rest empty,
    which is the case the quality gate exists for: a mostly-image document with
    a text cover page or an OCR'd title. A fully-empty PDF is caught earlier and
    more crudely by ``load_file``'s "no segments" check, so it does not exercise
    the density gate at all.
    """
    texts = ["Scanned document cover page."] * text_pages
    return make_pdf(texts + [""] * (page_count - text_pages))


def make_docx(sections: list[tuple[str, str]]) -> bytes:
    """Build a DOCX from (heading, body) pairs; headings use the Heading 1 style."""
    document = DocxDocument()
    for heading, body in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(body)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
